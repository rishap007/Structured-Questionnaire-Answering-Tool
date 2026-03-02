from __future__ import annotations
import io
import json
import logging
import os
import re
from pathlib import Path
from typing import Literal, Optional

import pandas as pd
from docx import Document
from fastapi import Depends, FastAPI, File, HTTPException, Query, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pypdf import PdfReader
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from sqlalchemy import inspect, text
from sqlalchemy.orm import Session

from .auth import get_current_user, router as auth_router
from .config import settings
from .database import Base, engine, get_db
from .models import (
    EditedAnswer,
    GeneratedAnswer,
    GenerationRun,
    Question,
    Questionnaire,
    ReferenceChunk,
    ReferenceDocument,
    RunAnswer,
    User,
)
from .rag.chunker import chunk_document
from .rag.generator import generator
from .rag.retriever import retriever
from .schemas import (
    EditAnswerRequest,
    CompareRunsOut,
    GenerateAnswersRequest,
    QuestionnaireOut,
    QuestionnaireResultOut,
    ReferenceDocumentOut,
    RunSummaryOut,
)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def ensure_schema_updates():
    inspector = inspect(engine)
    tables = set(inspector.get_table_names())

    if "generated_answers" in tables:
        cols = {c["name"] for c in inspector.get_columns("generated_answers")}
        with engine.begin() as conn:
            if "evidence_snippets" not in cols:
                conn.execute(text("ALTER TABLE generated_answers ADD COLUMN evidence_snippets TEXT DEFAULT '[]'"))


ensure_schema_updates()
Base.metadata.create_all(bind=engine)
os.makedirs(settings.upload_dir, exist_ok=True)
os.makedirs(settings.chroma_persist_dir, exist_ok=True)

app = FastAPI(title=settings.app_name)
app.include_router(auth_router)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def save_upload_file(file: UploadFile) -> str:
    file_path = Path(settings.upload_dir) / file.filename
    with file_path.open("wb") as out:
        out.write(file.file.read())
    return str(file_path)


def parse_pdf(file_path: str) -> str:
    reader = PdfReader(file_path)
    pages = []
    for p in reader.pages:
        pages.append(p.extract_text() or "")
    return "\n".join(pages)


def parse_text(file_path: str) -> str:
    return Path(file_path).read_text(encoding="utf-8", errors="ignore")


def parse_spreadsheet_questions(file_path: str) -> list[str]:
    questions: list[str] = []
    if file_path.endswith(".csv"):
        try:
            df = pd.read_csv(file_path)
        except Exception:
            df = pd.read_csv(file_path, engine="python", on_bad_lines="skip")
    else:
        df = pd.read_excel(file_path)

    for _, row in df.iterrows():
        for cell in row.tolist():
            if isinstance(cell, str):
                text = cell.strip()
                if len(text) > 8 and "?" in text:
                    questions.append(text)
                    break
    return questions


def extract_questions_from_text(text: str) -> list[str]:
    candidates: list[str] = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        line = re.sub(r"^(Q\d+[:.)\-\s]+|\d+[:.)\-\s]+)", "", line).strip()
        if len(line) > 8 and "?" in line:
            candidates.append(line)

    seen = set()
    unique = []
    for c in candidates:
        key = c.lower()
        if key not in seen:
            seen.add(key)
            unique.append(c)
    return unique


def load_document_text(file_path: str) -> str:
    if file_path.lower().endswith(".pdf"):
        return parse_pdf(file_path)
    if file_path.lower().endswith(".txt"):
        return parse_text(file_path)
    raise HTTPException(status_code=400, detail="Only PDF and TXT are supported for reference documents")


def load_questionnaire_questions(file_path: str) -> list[str]:
    lower = file_path.lower()
    if lower.endswith(".pdf"):
        return extract_questions_from_text(parse_pdf(file_path))
    if lower.endswith(".csv") or lower.endswith(".xlsx"):
        return parse_spreadsheet_questions(file_path)
    raise HTTPException(status_code=400, detail="Only PDF, CSV, and XLSX are supported for questionnaires")


def latest_edit(answer: GeneratedAnswer) -> Optional[EditedAnswer]:
    if not answer.edits:
        return None
    return sorted(answer.edits, key=lambda x: x.created_at)[-1]


def confidence_level(score: float) -> str:
    if score >= 75:
        return "High"
    if score >= 50:
        return "Medium"
    return "Low"


def build_evidence_snippets(hits: list[dict], limit: int = 3) -> list[dict]:
    snippets = []
    for hit in hits[:limit]:
        md = hit.get("metadata", {}) or {}
        citation = f"[{md.get('document_name', 'Reference Document')} – {md.get('section_title', 'General')}]"
        text_val = (hit.get("text") or "").strip()
        snippet_text = text_val[:280].strip()
        if len(text_val) > 280:
            snippet_text += "..."
        snippets.append(
            {
                "citation": citation,
                "snippet_text": snippet_text,
                "similarity": round(float(hit.get("similarity", 0.0)), 4),
            }
        )
    return snippets


def parse_evidence_snippets(raw: Optional[str]) -> list[dict]:
    if not raw:
        return []
    try:
        val = json.loads(raw)
    except Exception:
        return []
    return val if isinstance(val, list) else []


def run_summary(run: GenerationRun) -> dict:
    return {"id": run.id, "run_number": run.run_number, "created_at": run.created_at}


@app.get("/health")
def health():
    return {"status": "ok"}


@app.post("/reference-documents/upload", response_model=ReferenceDocumentOut)
def upload_reference_document(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    if not file.filename.lower().endswith((".pdf", ".txt")):
        raise HTTPException(status_code=400, detail="Reference file must be PDF or TXT")

    saved = save_upload_file(file)
    text = load_document_text(saved)
    if not text.strip():
        raise HTTPException(status_code=400, detail="Could not extract text from document")

    doc = ReferenceDocument(
        user_id=current_user.id,
        name=Path(file.filename).stem,
        source_filename=file.filename,
        content_type=file.content_type or "application/octet-stream",
        content_text=text,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    chunks = chunk_document(text, chunk_size=700, overlap=100)
    payload = []
    for c in chunks:
        vector_id = f"doc-{doc.id}-chunk-{c.chunk_index}"
        db_chunk = ReferenceChunk(
            document_id=doc.id,
            chunk_index=c.chunk_index,
            section_title=c.section_title,
            chunk_text=c.text,
            vector_id=vector_id,
        )
        db.add(db_chunk)
        payload.append(
            {
                "vector_id": vector_id,
                "document_id": str(doc.id),
                "document_name": doc.name,
                "section_title": c.section_title,
                "chunk_text": c.text,
            }
        )
    db.commit()
    retriever.add_chunks(payload)

    return doc


@app.get("/reference-documents", response_model=list[ReferenceDocumentOut])
def list_reference_documents(db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    return (
        db.query(ReferenceDocument)
        .filter(ReferenceDocument.user_id == current_user.id)
        .order_by(ReferenceDocument.created_at.desc())
        .all()
    )


@app.post("/questionnaires/upload", response_model=QuestionnaireOut)
def upload_questionnaire(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    if not file.filename.lower().endswith((".pdf", ".csv", ".xlsx")):
        raise HTTPException(status_code=400, detail="Questionnaire must be PDF, CSV, or XLSX")

    saved = save_upload_file(file)
    questions = load_questionnaire_questions(saved)
    if not questions:
        raise HTTPException(status_code=400, detail="No questions detected in questionnaire")

    questionnaire = Questionnaire(
        user_id=current_user.id,
        name=Path(file.filename).stem,
        source_filename=file.filename,
    )
    db.add(questionnaire)
    db.commit()
    db.refresh(questionnaire)

    for idx, q_text in enumerate(questions, start=1):
        q = Question(
            questionnaire_id=questionnaire.id,
            question_uid=f"Q-{idx:03d}",
            position=idx,
            text=q_text,
        )
        db.add(q)
    db.commit()

    return questionnaire


@app.get("/questionnaires", response_model=list[QuestionnaireOut])
def list_questionnaires(db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    return db.query(Questionnaire).filter(Questionnaire.user_id == current_user.id).order_by(Questionnaire.created_at.desc()).all()


@app.get("/questionnaires/{questionnaire_id}")
def questionnaire_detail(
    questionnaire_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    questionnaire = (
        db.query(Questionnaire)
        .filter(Questionnaire.id == questionnaire_id, Questionnaire.user_id == current_user.id)
        .first()
    )
    if not questionnaire:
        raise HTTPException(status_code=404, detail="Questionnaire not found")

    questions = (
        db.query(Question)
        .filter(Question.questionnaire_id == questionnaire.id)
        .order_by(Question.position.asc())
        .all()
    )
    return {
        "questionnaire": QuestionnaireOut.model_validate(questionnaire),
        "questions": [
            {"id": q.id, "question_uid": q.question_uid, "position": q.position, "text": q.text}
            for q in questions
        ],
    }


@app.post("/questionnaires/{questionnaire_id}/generate")
def generate_answers(
    questionnaire_id: int,
    payload: Optional[GenerateAnswersRequest] = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    questionnaire = (
        db.query(Questionnaire)
        .filter(Questionnaire.id == questionnaire_id, Questionnaire.user_id == current_user.id)
        .first()
    )
    if not questionnaire:
        raise HTTPException(status_code=404, detail="Questionnaire not found")

    questions = db.query(Question).filter(Question.questionnaire_id == questionnaire.id).order_by(Question.position.asc()).all()
    if not questions:
        raise HTTPException(status_code=400, detail="No questions found")

    user_doc_ids = [
        str(doc_id)
        for (doc_id,) in db.query(ReferenceDocument.id).filter(ReferenceDocument.user_id == current_user.id).all()
    ]

    selected_ids = set(payload.question_ids if payload else [])
    selected_uids = {u.strip() for u in (payload.question_uids if payload else []) if u.strip()}
    selected_question_ids = {
        q.id for q in questions if (not selected_ids and not selected_uids) or q.id in selected_ids or q.question_uid in selected_uids
    }
    if not selected_question_ids:
        raise HTTPException(status_code=400, detail="No valid selected questions for regeneration")

    latest_run = (
        db.query(GenerationRun)
        .filter(GenerationRun.questionnaire_id == questionnaire.id)
        .order_by(GenerationRun.run_number.desc())
        .first()
    )
    next_run_number = (latest_run.run_number + 1) if latest_run else 1
    run = GenerationRun(questionnaire_id=questionnaire.id, run_number=next_run_number)
    db.add(run)
    db.commit()
    db.refresh(run)

    for question in questions:
        existing = db.query(GeneratedAnswer).filter(GeneratedAnswer.question_id == question.id).first()

        if question.id in selected_question_ids:
            hits = retriever.retrieve(question.text, k=3, document_ids=user_doc_ids) if user_doc_ids else []
            similarities = [h["similarity"] for h in hits]
            max_similarity = max(similarities, default=0.0)
            avg_similarity = sum(similarities) / len(similarities) if similarities else 0.0
            logger.info("Question %s similarities=%s", question.question_uid, similarities)

            if max_similarity < settings.rag_similarity_threshold:
                answer_text = "Not found in references."
                citations = ""
                evidence_snippets = []
            else:
                answer_text, citations = generator.generate(question.text, hits)
                evidence_snippets = build_evidence_snippets(hits)
                if answer_text != "Not found in references." and not citations.strip():
                    answer_text = "Not found in references."

            confidence = round(avg_similarity * 100, 2)
        else:
            if existing:
                answer_text = existing.answer_text
                citations = existing.citations
                evidence_snippets = parse_evidence_snippets(existing.evidence_snippets)
                confidence = existing.confidence_score
                avg_similarity = existing.avg_similarity_score
            else:
                answer_text = "Not found in references."
                citations = ""
                evidence_snippets = []
                confidence = 0.0
                avg_similarity = 0.0

        evidence_json = json.dumps(evidence_snippets, ensure_ascii=True)
        if existing:
            existing.answer_text = answer_text
            existing.citations = citations
            existing.evidence_snippets = evidence_json
            existing.confidence_score = confidence
            existing.avg_similarity_score = avg_similarity
        else:
            db.add(
                GeneratedAnswer(
                    question_id=question.id,
                    answer_text=answer_text,
                    citations=citations,
                    evidence_snippets=evidence_json,
                    confidence_score=confidence,
                    avg_similarity_score=avg_similarity,
                )
            )
        db.add(
            RunAnswer(
                run_id=run.id,
                question_id=question.id,
                answer_text=answer_text,
                citations=citations,
                evidence_snippets=evidence_json,
                confidence_score=confidence,
                avg_similarity_score=avg_similarity,
            )
        )

    db.commit()
    return {
        "message": "Answers generated successfully",
        "run_id": run.id,
        "run_number": run.run_number,
        "regenerated_questions": len(selected_question_ids),
        "total_questions": len(questions),
    }


def build_results(questionnaire: Questionnaire, db: Session, run_id: Optional[int] = None) -> dict:
    questions = db.query(Question).filter(Question.questionnaire_id == questionnaire.id).order_by(Question.position.asc()).all()
    latest_run = (
        db.query(GenerationRun)
        .filter(GenerationRun.questionnaire_id == questionnaire.id)
        .order_by(GenerationRun.run_number.desc())
        .first()
    )
    target_run = None
    if run_id is not None:
        target_run = (
            db.query(GenerationRun)
            .filter(GenerationRun.id == run_id, GenerationRun.questionnaire_id == questionnaire.id)
            .first()
        )
        if not target_run:
            raise HTTPException(status_code=404, detail="Run not found")
    else:
        target_run = latest_run

    run_answer_map = {}
    if target_run:
        run_answers = db.query(RunAnswer).filter(RunAnswer.run_id == target_run.id).all()
        run_answer_map = {ra.question_id: ra for ra in run_answers}

    answer_rows = []
    answered_with_citations = 0
    not_found_count = 0

    for q in questions:
        g = db.query(GeneratedAnswer).filter(GeneratedAnswer.question_id == q.id).first()
        run_answer = run_answer_map.get(q.id)
        if not g and not run_answer:
            answer_text = "Not found in references."
            citations = ""
            confidence_score = 0.0
            avg_similarity_score = 0.0
            evidence_snippets = []
            generated_answer_id = 0
            edit = None
        else:
            generated_answer_id = g.id if g else 0
            answer_text = run_answer.answer_text if run_answer else g.answer_text
            citations = run_answer.citations if run_answer else g.citations
            confidence_score = run_answer.confidence_score if run_answer else g.confidence_score
            avg_similarity_score = run_answer.avg_similarity_score if run_answer else g.avg_similarity_score
            evidence_snippets = parse_evidence_snippets(run_answer.evidence_snippets if run_answer else g.evidence_snippets)
            # Apply manual edits only in current/latest view.
            edit = latest_edit(g) if g and (run_id is None or (latest_run and target_run and latest_run.id == target_run.id)) else None

        # Review flow always displays latest edit if present, while preserving generated baseline.
        effective_text = edit.edited_text if edit else answer_text
        effective_citations = edit.edited_citations if edit else citations

        if effective_text.strip() == "Not found in references.":
            not_found_count += 1
        elif effective_citations.strip():
            answered_with_citations += 1

        answer_rows.append(
            {
                "generated_answer_id": generated_answer_id,
                "question_id": q.id,
                "question_uid": q.question_uid,
                "position": q.position,
                "question_text": q.text,
                "generated_answer_text": answer_text,
                "generated_citations": citations,
                "edited_answer_text": edit.edited_text if edit else None,
                "edited_citations": edit.edited_citations if edit else None,
                "evidence_snippets": evidence_snippets,
                "confidence_score": confidence_score,
                "confidence_level": confidence_level(confidence_score),
                "avg_similarity_score": avg_similarity_score,
            }
        )

    return {
        "questionnaire": QuestionnaireOut.model_validate(questionnaire),
        "run": run_summary(target_run) if target_run else None,
        "summary": {
            "total_questions": len(questions),
            "answered_with_citations": answered_with_citations,
            "not_found_count": not_found_count,
        },
        "answers": answer_rows,
    }


@app.get("/questionnaires/{questionnaire_id}/results", response_model=QuestionnaireResultOut)
def questionnaire_results(
    questionnaire_id: int,
    run_id: Optional[int] = Query(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    questionnaire = (
        db.query(Questionnaire)
        .filter(Questionnaire.id == questionnaire_id, Questionnaire.user_id == current_user.id)
        .first()
    )
    if not questionnaire:
        raise HTTPException(status_code=404, detail="Questionnaire not found")
    return build_results(questionnaire, db, run_id=run_id)


@app.get("/questionnaires/{questionnaire_id}/runs", response_model=list[RunSummaryOut])
def list_questionnaire_runs(
    questionnaire_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    questionnaire = (
        db.query(Questionnaire)
        .filter(Questionnaire.id == questionnaire_id, Questionnaire.user_id == current_user.id)
        .first()
    )
    if not questionnaire:
        raise HTTPException(status_code=404, detail="Questionnaire not found")
    runs = (
        db.query(GenerationRun)
        .filter(GenerationRun.questionnaire_id == questionnaire.id)
        .order_by(GenerationRun.run_number.desc())
        .all()
    )
    return [run_summary(r) for r in runs]


@app.get("/questionnaires/{questionnaire_id}/runs/compare", response_model=CompareRunsOut)
def compare_questionnaire_runs(
    questionnaire_id: int,
    old_run_id: int = Query(...),
    new_run_id: int = Query(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    questionnaire = (
        db.query(Questionnaire)
        .filter(Questionnaire.id == questionnaire_id, Questionnaire.user_id == current_user.id)
        .first()
    )
    if not questionnaire:
        raise HTTPException(status_code=404, detail="Questionnaire not found")

    old_run = (
        db.query(GenerationRun)
        .filter(GenerationRun.id == old_run_id, GenerationRun.questionnaire_id == questionnaire.id)
        .first()
    )
    new_run = (
        db.query(GenerationRun)
        .filter(GenerationRun.id == new_run_id, GenerationRun.questionnaire_id == questionnaire.id)
        .first()
    )
    if not old_run or not new_run:
        raise HTTPException(status_code=404, detail="One or more runs not found")

    old_map = {a.question_id: a for a in db.query(RunAnswer).filter(RunAnswer.run_id == old_run.id).all()}
    new_map = {a.question_id: a for a in db.query(RunAnswer).filter(RunAnswer.run_id == new_run.id).all()}
    questions = db.query(Question).filter(Question.questionnaire_id == questionnaire.id).order_by(Question.position.asc()).all()

    items = []
    changed_answers = 0
    for q in questions:
        old_val = old_map.get(q.id)
        new_val = new_map.get(q.id)
        old_text = old_val.answer_text if old_val else ""
        new_text = new_val.answer_text if new_val else ""
        old_conf = old_val.confidence_score if old_val else 0.0
        new_conf = new_val.confidence_score if new_val else 0.0
        changed = old_text.strip() != new_text.strip()
        if changed:
            changed_answers += 1
        items.append(
            {
                "question_id": q.id,
                "question_uid": q.question_uid,
                "position": q.position,
                "question_text": q.text,
                "old_answer_text": old_text,
                "new_answer_text": new_text,
                "old_confidence_score": old_conf,
                "new_confidence_score": new_conf,
                "changed": changed,
            }
        )

    return {
        "questionnaire": QuestionnaireOut.model_validate(questionnaire),
        "old_run": run_summary(old_run),
        "new_run": run_summary(new_run),
        "total_questions": len(questions),
        "changed_answers": changed_answers,
        "items": items,
    }


@app.put("/answers/{generated_answer_id}/edit")
def edit_answer(
    generated_answer_id: int,
    payload: EditAnswerRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    generated = db.query(GeneratedAnswer).filter(GeneratedAnswer.id == generated_answer_id).first()
    if not generated:
        raise HTTPException(status_code=404, detail="Answer not found")

    q = db.query(Question).filter(Question.id == generated.question_id).first()
    questionnaire = db.query(Questionnaire).filter(Questionnaire.id == q.questionnaire_id).first()
    if questionnaire.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Forbidden")

    edit = EditedAnswer(
        generated_answer_id=generated_answer_id,
        edited_text=payload.edited_text,
        edited_citations=payload.edited_citations,
    )
    db.add(edit)
    db.commit()
    return {"message": "Edit saved"}


def export_docx(results: dict) -> bytes:
    document = Document()
    document.add_heading(results["questionnaire"].name, level=1)

    for item in results["answers"]:
        answer = item["edited_answer_text"] or item["generated_answer_text"]
        citations = item["edited_citations"] or item["generated_citations"]
        document.add_paragraph(f"{item['position']}. {item['question_text']}", style="Heading 3")
        document.add_paragraph(answer)
        if citations:
            document.add_paragraph(citations)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()


def export_pdf(results: dict) -> bytes:
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    y = height - 40

    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(40, y, results["questionnaire"].name)
    y -= 28

    for item in results["answers"]:
        answer = item["edited_answer_text"] or item["generated_answer_text"]
        citations = item["edited_citations"] or item["generated_citations"]

        lines = [f"{item['position']}. {item['question_text']}", f"Answer: {answer}"]
        if citations:
            lines.append(f"Citations: {citations}")

        for line in lines:
            wrapped = [line[i : i + 105] for i in range(0, len(line), 105)]
            for piece in wrapped:
                if y < 60:
                    pdf.showPage()
                    y = height - 40
                    pdf.setFont("Helvetica", 10)
                pdf.drawString(40, y, piece)
                y -= 14
        y -= 8

    pdf.save()
    buffer.seek(0)
    return buffer.read()


@app.get("/questionnaires/{questionnaire_id}/export")
def export_questionnaire(
    questionnaire_id: int,
    format: Literal["pdf", "docx"] = Query("docx"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    questionnaire = (
        db.query(Questionnaire)
        .filter(Questionnaire.id == questionnaire_id, Questionnaire.user_id == current_user.id)
        .first()
    )
    if not questionnaire:
        raise HTTPException(status_code=404, detail="Questionnaire not found")

    results = build_results(questionnaire, db)

    if format == "docx":
        content = export_docx(results)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ext = "docx"
    else:
        content = export_pdf(results)
        media_type = "application/pdf"
        ext = "pdf"

    filename = f"{questionnaire.name}_answered.{ext}"
    return StreamingResponse(
        io.BytesIO(content),
        media_type=media_type,
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )
